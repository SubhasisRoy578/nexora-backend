# ==================================================
# NEXORA AI — FILE READER
# Supports: PDF, DOCX, TXT, CSV, Excel, Markdown
# Each returns: { text, metadata, pages, file_type }
# ==================================================

import os
import io
from pathlib import Path


# ==================================================
# PDF
# ==================================================

def read_pdf(path: str) -> dict:

    try:

        from pypdf import PdfReader

        reader = PdfReader(path)

        pages = []
        full_text = ""

        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            pages.append({
                "page": i + 1,
                "text": page_text
            })
            full_text += page_text + "\n"

        return {
            "success": True,
            "file_type": "pdf",
            "text": full_text.strip(),
            "pages": pages,
            "page_count": len(reader.pages),
            "metadata": dict(reader.metadata) if reader.metadata else {}
        }

    except Exception as e:

        return {
            "success": False,
            "file_type": "pdf",
            "text": "",
            "error": str(e)
        }


# ==================================================
# DOCX
# ==================================================

def read_docx(path: str) -> dict:

    try:

        import docx

        doc = docx.Document(path)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n".join(paragraphs)

        # Extract tables
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(rows)

        return {
            "success": True,
            "file_type": "docx",
            "text": full_text,
            "paragraphs": paragraphs,
            "tables": tables,
            "paragraph_count": len(paragraphs),
            "metadata": {}
        }

    except Exception as e:

        return {
            "success": False,
            "file_type": "docx",
            "text": "",
            "error": str(e)
        }


# ==================================================
# TXT / MARKDOWN
# ==================================================

def read_txt(path: str) -> dict:

    try:

        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()

        ext = Path(path).suffix.lower()

        return {
            "success": True,
            "file_type": ext.replace(".", "") or "txt",
            "text": text,
            "char_count": len(text),
            "line_count": len(text.splitlines()),
            "metadata": {}
        }

    except Exception as e:

        return {
            "success": False,
            "file_type": "txt",
            "text": "",
            "error": str(e)
        }


# ==================================================
# CSV
# ==================================================

def read_csv(path: str) -> dict:

    try:

        import csv

        rows = []
        headers = []

        with open(path, "r", encoding="utf-8", errors="ignore") as f:

            reader = csv.DictReader(f)
            headers = reader.fieldnames or []

            for row in reader:
                rows.append(dict(row))

        # Build readable text summary
        preview_rows = rows[:10]
        text_lines = [",".join(headers)]
        for row in preview_rows:
            text_lines.append(",".join(str(v) for v in row.values()))

        full_text = "\n".join(text_lines)

        return {
            "success": True,
            "file_type": "csv",
            "text": full_text,
            "headers": headers,
            "rows": rows,
            "row_count": len(rows),
            "preview": preview_rows,
            "metadata": {}
        }

    except Exception as e:

        return {
            "success": False,
            "file_type": "csv",
            "text": "",
            "error": str(e)
        }


# ==================================================
# EXCEL (.xlsx / .xls)
# ==================================================

def read_excel(path: str) -> dict:

    try:

        import openpyxl

        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)

        sheets = {}
        full_text = ""

        for sheet_name in wb.sheetnames:

            ws = wb[sheet_name]
            rows = []

            for row in ws.iter_rows(values_only=True):
                cleaned = [str(cell) if cell is not None else "" for cell in row]
                rows.append(cleaned)

            sheets[sheet_name] = rows

            # Add sheet text to full_text
            full_text += f"\n[Sheet: {sheet_name}]\n"
            for row in rows[:20]:
                full_text += ",".join(row) + "\n"

        wb.close()

        return {
            "success": True,
            "file_type": "excel",
            "text": full_text.strip(),
            "sheets": sheets,
            "sheet_names": wb.sheetnames,
            "metadata": {}
        }

    except Exception as e:

        return {
            "success": False,
            "file_type": "excel",
            "text": "",
            "error": str(e)
        }


# ==================================================
# IMAGE — OCR via pytesseract
# ==================================================

def read_image(path: str) -> dict:

    try:

        import pytesseract
        from PIL import Image

        img = Image.open(path)
        text = pytesseract.image_to_string(img)

        return {
            "success": True,
            "file_type": "image",
            "text": text.strip(),
            "width": img.width,
            "height": img.height,
            "mode": img.mode,
            "metadata": {}
        }

    except Exception as e:

        return {
            "success": False,
            "file_type": "image",
            "text": "",
            "error": str(e)
        }


# ==================================================
# UNIVERSAL READER
# Auto-detects file type from extension
# ==================================================

def read_file(path: str) -> dict:
    """
    Auto-detects and reads any supported file.
    Returns unified dict: { success, file_type, text, metadata, ... }
    """

    ext = Path(path).suffix.lower()

    readers = {
        ".pdf":  read_pdf,
        ".docx": read_docx,
        ".doc":  read_docx,
        ".txt":  read_txt,
        ".md":   read_txt,
        ".csv":  read_csv,
        ".xlsx": read_excel,
        ".xls":  read_excel,
        ".png":  read_image,
        ".jpg":  read_image,
        ".jpeg": read_image,
        ".webp": read_image,
    }

    reader_fn = readers.get(ext)

    if not reader_fn:
        return {
            "success": False,
            "file_type": ext.replace(".", "") or "unknown",
            "text": "",
            "error": f"Unsupported file type: {ext}"
        }

    return reader_fn(path)


# ==================================================
# CHUNK TEXT FOR RAG
# Splits large text into overlapping chunks
# ==================================================

def chunk_text(
    text: str,
    chunk_size: int = 500,
    overlap: int = 50
) -> list:
    """
    Splits text into overlapping chunks for RAG ingestion.
    Returns list of strings.
    """

    if not text:
        return []

    words = text.split()
    chunks = []
    start = 0

    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        chunks.append(chunk)
        start += chunk_size - overlap

    return chunks